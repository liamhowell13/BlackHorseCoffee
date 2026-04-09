from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# --- Helper functions ---
def add_title(text):
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_page_header(text):
    doc.add_page_break()
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x2D, 0x5A, 0x27)

def add_section_header(text):
    doc.add_heading(text, level=2)

def add_subsection(text):
    doc.add_heading(text, level=3)

def add_field(label, value):
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(value)

def add_body(text):
    doc.add_paragraph(text)

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
    doc.add_paragraph()  # spacing

# ============================================================
# TITLE & INSTRUCTIONS
# ============================================================
add_title("BlackHorse Espresso\nWebsite Copy Review")

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Instructions for the reviewer:")
run.bold = True
run.font.size = Pt(12)

instructions = [
    "Review all text below. If something needs to change, just replace it directly with the corrected version.",
    "If something looks good, leave it as-is.",
    "Each section is labeled with the page and area so we know exactly where it appears on the website.",
    "Do NOT change the formatting or structure (headings, bullet points, tables) \u2014 just the words themselves.",
    "Pay special attention to: menu items, prices, hours, phone numbers, addresses, and descriptions.",
]
for inst in instructions:
    doc.add_paragraph(inst, style='List Bullet')

doc.add_paragraph()

# ============================================================
# GLOBAL ELEMENTS
# ============================================================
add_page_header("GLOBAL ELEMENTS")

add_section_header("Navigation Bar (appears on every page)")
add_body("Home  |  Menu  |  Locations  |  About  |  Silver Card  |  Contact")

add_section_header("Footer (appears on every page)")
add_field("Copyright", "\u00a9 2025 BlackHorse Espresso. All rights reserved.")
add_field("Footer links", "Menu  |  Locations  |  Contact")
add_field("Instagram", "@blackhorse.espresso")

add_section_header("Browser Tab / SEO")
add_field("Browser tab title", "BlackHorse Espresso")
add_field("Meta description", "Handcrafted coffee and fresh pastries in San Luis Obispo. Visit our two locations for Caribbean Coffee Company\u2019s medium dark Brazilian blend and artisan baked goods.")

# ============================================================
# HOME PAGE
# ============================================================
add_page_header("HOME PAGE")

add_section_header("Hero Banner")
add_field("Small text", "Est. 1995 \u00b7 San Luis Obispo")
add_field("Main heading", "BlackHorse")
add_field("Subheading", "Espresso & Bakery")
add_field("Description", "Handcrafted coffee and fresh pastries, made with care every single day.")
add_field("Button 1", "Visit Us")
add_field("Button 2", "View Menu")

add_section_header("About Preview Section")
add_field("Label", "Our Story")
add_field("Heading", "Coffee & Community Go Together")
add_field("Paragraph 1", "BlackHorse was born out of an idea that coffee and community go together. What started over 30 years ago has grown into two beloved locations in San Luis Obispo, each with its own personality but sharing the same commitment to quality.")
add_field("Paragraph 2", "We offer a selection of fresh pastries and baked goods daily. Paired with Caribbean Coffee Company\u2019s signature blend and hand-crafted espresso drinks, every visit is a treat for the senses.")

add_subsection("Feature Cards")
add_table(
    ["Card Title", "Card Text"],
    [
        ["Local Roasters", "We proudly serve Caribbean Coffee Company\u2019s medium dark Brazilian blend for a rich, bold cup every time."],
        ["Handcrafted Drinks", "Every drink is made to order by our skilled baristas who are passionate about the craft of coffee."],
        ["30+ Years", "Serving the San Luis Obispo community since 1995 with quality coffee and fresh pastries daily."],
    ]
)

add_section_header("Quote Section")
add_field("Quote", "\u201cCoffee should be black as hell, strong as death and as sweet as love.\u201d")
add_field("Attribution", "Turkish Proverb")

add_section_header("Gallery Section")
add_field("Label", "Made Fresh Daily")
add_field("Heading", "From Our Kitchen & Bar")

add_section_header("Locations Preview Section")
add_field("Label", "Come Say Hello")
add_field("Heading", "Our Locations")
add_field("Button", "View All Locations")

# ============================================================
# ABOUT PAGE
# ============================================================
add_page_header("ABOUT PAGE")

add_section_header("Hero Banner")
add_field("Label", "Est. 1995")
add_field("Heading", "Our Story")
add_field("Subheading", "Born from a simple idea: coffee and community go together.")

add_section_header("Our Beginning Section")
add_field("Label", "Our Beginning")
add_field("Heading", "From a Dream to Two Doors Open")
add_field("Paragraph 1", "BlackHorse was born out of an idea that coffee and community go together. What began over 30 years ago in downtown San Luis Obispo has grown into two beloved locations, each with its own character but sharing the same heart.")
add_field("Paragraph 2", "We proudly serve Caribbean Coffee Company\u2019s medium dark Brazilian blend. Our talented baristas craft every drink by hand, while our selection of fresh pastries and baked goods complements every cup.")

add_section_header("Our Values Section")
add_field("Label", "Our Values")
add_field("Heading", "What We Stand For")
add_table(
    ["Value", "Description"],
    [
        ["Quality First", "We source our beans from Caribbean Coffee Company, a rich medium dark Brazilian blend. Every cup is crafted with care and precision."],
        ["Community Driven", "BlackHorse is more than a coffee shop \u2014 it\u2019s a gathering place. We believe in building connections, one cup at a time."],
        ["Fresh & Local", "We offer fresh pastries and baked goods daily using quality ingredients. We partner with local suppliers whenever possible."],
        ["Craft & Passion", "Our baristas are artisans. From latte art to the perfect pour, we take pride in mastering our craft."],
    ]
)

add_section_header("Timeline Section")
add_field("Label", "Milestones")
add_field("Heading", "Our Journey")
add_table(
    ["Year", "Event"],
    [
        ["1995", "BlackHorse Espresso opens its first location on Higuera St. in downtown SLO."],
        ["2011", "Second location opens on Broad St., bringing BlackHorse to more of the community."],
        ["2025", "Celebrating 30 years of serving San Luis Obispo with quality coffee and community."],
    ]
)

add_section_header("Quote Section")
add_field("Quote", "\u201cCoffee should be black as hell, strong as death and as sweet as love.\u201d")
add_field("Attribution", "Turkish Proverb")

# ============================================================
# MENU PAGE
# ============================================================
add_page_header("MENU PAGE")

add_section_header("Hero Banner")
add_field("Label", "BlackHorse Espresso")
add_field("Heading", "Our Menu")

add_section_header("HOT DRINKS")

add_subsection("Espresso")
add_table(
    ["Item", "8oz", "12oz", "16oz", "20oz"],
    [
        ["Espresso", "$3.50", "\u2013", "\u2013", "\u2013"],
        ["Macchiato", "$4.25", "\u2013", "\u2013", "\u2013"],
        ["Espresso con Panna", "$4.25", "\u2013", "\u2013", "\u2013"],
    ]
)

add_subsection("Coffee & Lattes")
add_table(
    ["Item", "8oz", "12oz", "16oz", "20oz"],
    [
        ["Coffee of the Day", "$3.75", "$4.00", "$4.25", "$4.50"],
        ["Caffe Americano", "\u2013", "$4.50", "$5.15", "$5.75"],
        ["Au Lait", "$4.00", "$4.25", "$4.50", "$4.75"],
        ["Cappuccino", "$4.75", "$5.25", "$5.75", "$6.50"],
        ["Caffe Latte", "$4.75", "$5.25", "$5.75", "$6.50"],
        ["Caffe Mocha", "$5.25", "$5.80", "$6.45", "$7.25"],
        ["White Mocha", "$5.25", "$5.80", "$6.45", "$7.25"],
        ["Caramello", "$5.25", "$5.80", "$6.45", "$7.25"],
        ["Caffe Generra", "$5.25", "$5.80", "$6.45", "$7.25"],
        ["Matcha Latte", "$5.25", "$5.80", "$6.45", "$7.25"],
    ]
)

add_subsection("Other Hot Drinks")
add_table(
    ["Item", "8oz", "12oz", "16oz", "20oz"],
    [
        ["Cocoa", "$3.90", "$4.30", "$4.70", "$5.10"],
        ["Steamed Moo", "$3.15", "$3.55", "$3.95", "$4.35"],
        ["Mulled Cider", "$3.95", "$4.35", "$4.75", "$5.15"],
        ["Chai Tea", "$4.95", "$5.35", "$5.75", "$6.50"],
        ["Select Teas", "\u2013", "$3.25", "$3.50", "\u2013"],
    ]
)

add_section_header("COLD DRINKS")

add_subsection("Iced Coffee")
add_table(
    ["Item", "Description", "12oz", "16oz", "20oz"],
    [
        ["Iced Americano", "", "$4.50", "$5.15", "$5.75"],
        ["Iced Latte", "", "$5.15", "$5.75", "$6.50"],
        ["Iced Matcha", "", "$5.15", "$5.75", "$6.50"],
        ["Iced Tea", "Paradise passion fruit", "$3.50", "$3.75", "$4.25"],
        ["Fresh Fruit Juice", "", "$3.60", "$4.00", "$4.40"],
    ]
)

add_subsection("Blended & Cold")
add_table(
    ["Item", "12oz", "16oz", "20oz"],
    [
        ["Cold Milk", "$3.05", "$3.45", "$3.85"],
        ["Uptown Squeeze", "\u2013", "$6.00", "$6.50"],
        ["Uptown Freeze", "\u2013", "$6.00", "$7.25"],
        ["Flavored Freezes", "\u2013", "$7.00", "$7.50"],
    ]
)

add_section_header("BAKED GOODS")

add_subsection("Sensational Treats")
add_table(
    ["Item", "Description", "Price"],
    [
        ["Butter Croissant", "Flaky, laminated French-style croissant", "$4.25"],
        ["Almond Croissant", "Filled with almond cream, topped with sliced almonds", "$4.75"],
        ["Chocolate Croissant", "Dark chocolate batons in buttery pastry", "$4.75"],
        ["Morning Bun", "Cinnamon, sugar-coated laminated dough", "$4.50"],
        ["Blueberry Muffin", "Loaded with fresh blueberries, streusel top", "$4.00"],
        ["Banana Nut Bread", "Moist banana bread with walnuts", "$4.00"],
    ]
)

add_subsection("Savory")
add_table(
    ["Item", "Description", "Price"],
    [
        ["Ham & Cheese Croissant", "Black forest ham, gruyere, dijon", "$6.50"],
        ["Spinach Feta Croissant", "Spinach, feta, egg, sun-dried tomato", "$7.00"],
        ["Breakfast Burrito", "Scrambled eggs, cheese, potatoes, salsa", "$8.00"],
    ]
)

add_subsection("Cookies & Bars")
add_table(
    ["Item", "Description", "Price"],
    [
        ["Chocolate Chip Cookie", "", "$3.50"],
        ["Oatmeal Raisin Cookie", "", "$3.50"],
        ["Brownie", "Rich, fudgy dark chocolate", "$4.00"],
        ["Lemon Bar", "Buttery shortbread with tangy lemon curd", "$4.00"],
    ]
)

# ============================================================
# LOCATIONS PAGE
# ============================================================
add_page_header("LOCATIONS PAGE")

add_section_header("Hero Banner")
add_field("Label", "Come Say Hello")
add_field("Heading", "Our Locations")

add_section_header("Uptown Location")
add_field("Name", "Uptown")
add_field("Address", "1065 Higuera St.")
add_field("City", "San Luis Obispo, CA 93401")
add_field("Phone", "805-783-1300")
add_field("Hours", "Open 6:30 am \u2013 5:00 pm daily")
add_field("Amenities", "Full espresso bar, Fresh baked goods, Indoor & patio seating, Free Wi-Fi")

add_section_header("Broad Location")
add_field("Name", "Broad")
add_field("Address", "3590 Broad St.")
add_field("City", "San Luis Obispo, CA 93401")
add_field("Phone", "805-439-1300")
add_field("Hours", "Open 6:30 am \u2013 5:00 pm daily")
add_field("Amenities", "Full espresso bar, Fresh baked goods, Drive-through friendly area, Free Wi-Fi")

# ============================================================
# SILVER CARD PAGE
# ============================================================
add_page_header("SILVER CARD PAGE")

add_section_header("Hero Banner")
add_field("Heading", "Silver Card")
add_field("Subheading", "Get an extra boost with every purchase")

add_section_header("Intro Section")
add_field("Label", "Loyalty Program")
add_field("Heading", "Save 10% Off Every Visit")
add_field("Paragraph 1", "For our loyal customers who drop in often, the BlackHorse Silver Card saves you time and money! The Silver Card is a prepaid card that gets you 10% off on all your BlackHorse purchases all year-round.")
add_field("Paragraph 2", "Whether you\u2019re a daily regular or just love treating yourself, the Silver Card is the easiest way to make every cup count.")
add_field("Button", "Get Yours at Any Location")

add_subsection("Card Visual Text")
add_field("Top label", "Prepaid Loyalty")
add_field("Card title", "Silver Card")
add_field("Company name", "BlackHorse Espresso")
add_field("Discount text", "10% off every purchase")

add_section_header("Benefits Section")
add_field("Label", "Benefits")
add_field("Heading", "Why You\u2019ll Love It")
add_table(
    ["Benefit", "Description"],
    [
        ["10% Off Everything", "Save 10% on all your BlackHorse purchases all year-round \u2014 coffee, pastries, and more."],
        ["Access by Name", "Your account is accessible by name at either location. No need to carry the physical card."],
        ["Perfect for Gifting", "Great as gifts for employees, clients, friends, or anyone who loves quality coffee and baked goods."],
        ["Prepaid & Simple", "Load your card with any amount and enjoy the discount on every purchase until the balance runs out. Reload anytime."],
    ]
)

add_section_header("How It Works Section")
add_field("Label", "Getting Started")
add_field("Heading", "How It Works")
add_field("Subtext", "Sign up for 10% off today. It\u2019s that easy.")
add_table(
    ["Step", "Title", "Description"],
    [
        ["01", "Visit Any Location", "Stop by either our Uptown or Broad St. location and ask about the Silver Card."],
        ["02", "Load Your Card", "Add any amount to your prepaid Silver Card. The more you load, the more you save."],
        ["03", "Save 10% Every Time", "Enjoy 10% off on every purchase, every visit. It\u2019s that easy."],
    ]
)

add_section_header("Call-to-Action Section")
add_field("Heading", "Ready to Start Saving?")
add_field("Text", "Visit either of our locations to pick up your Silver Card and start enjoying 10% off every purchase.")
add_field("Button 1", "Find a Location")
add_field("Button 2", "Questions? Contact Us")

# ============================================================
# CONTACT PAGE
# ============================================================
add_page_header("CONTACT PAGE")

add_section_header("Hero Banner")
add_field("Label", "Get in Touch")
add_field("Heading", "Contact")

add_section_header("Contact Form")
add_field("Label", "Message Us")
add_field("Heading", "Send us a Message")
add_field("Form fields", "Name, Email, Subject, Message")
add_field("Button", "Send Message")

add_subsection("After Submission")
add_field("Heading", "Message Sent!")
add_field("Text", "Thank you for reaching out. We\u2019ll get back to you as soon as possible.")

add_section_header("Sidebar \u2014 Social")
add_field("Heading", "Follow Us")
add_field("Instagram", "@blackhorse.espresso")

# ============================================================
# SAVE
# ============================================================
doc.save("/Users/liamhowell/Documents/Code/Coffee/BlackHorse_Website_Copy_Review.docx")
print("Done!")
